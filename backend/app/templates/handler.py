import os
import json
from typing import Dict, List, Optional, Any
from dataclasses import dataclass
from docx import Document
from docx2pdf import convert
import platform
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.application import MIMEApplication
import smtplib
import tempfile


@dataclass
class Placeholder:
    """Class for storing placeholder information"""
    id: str
    placeholder_tag: str
    original_text: str
    type: str
    required: bool
    description: Dict[str, str]
    source: str


@dataclass
class TemplateInfo:
    """Class for storing template information"""
    id: str
    filename: str
    display_name: Dict[str, str]
    description: Dict[str, str]
    placeholders: List[Placeholder]


class TemplateHandler:
    def __init__(self, templates_dir: str = "app/templates/docs", templates_config_path: str = "app/templates/templates_config.json"):
        """Initialize the template handler with the templates directory and config file"""
        self.templates_dir = templates_dir
        self.templates_config_path = templates_config_path
        self.templates_config = self._load_templates_config()
        
        # Initialize COM for Windows if needed
        if platform.system() == 'Windows':
            try:
                import pythoncom
                self.pythoncom = pythoncom
            except ImportError:
                print("Warning: pythoncom not available. PDF conversion may not work on Windows.")
                self.pythoncom = None
        else:
            self.pythoncom = None
            
    def _load_templates_config(self) -> Dict[str, Any]:
        """Load the templates configuration from JSON file"""
        try:
            if os.path.exists(self.templates_config_path):
                with open(self.templates_config_path, 'r', encoding='utf-8') as f:
                    return json.load(f)
            else:
                print(f"Template configuration file not found at {self.templates_config_path}")
                return {"templates": []}
        except Exception as e:
            print(f"Error loading templates configuration: {str(e)}")
            return {"templates": []}
    
    def get_available_templates(self) -> List[TemplateInfo]:
        """Get all available templates and their placeholders from the config file"""
        templates = []
        
        for template_data in self.templates_config.get("templates", []):
            placeholders = []
            
            for placeholder_data in template_data.get("placeholders", []):
                # Create Placeholder object
                placeholder = Placeholder(
                    id=placeholder_data.get("id", ""),
                    placeholder_tag=placeholder_data.get("placeholder_tag", ""),
                    original_text=placeholder_data.get("original_text", ""),
                    type=placeholder_data.get("type", "string"),
                    required=placeholder_data.get("required", True),
                    description=placeholder_data.get("description", {"ar": "", "en": ""}),
                    source=placeholder_data.get("source", "user_input")
                )
                placeholders.append(placeholder)
            
            # Create TemplateInfo object
            template_info = TemplateInfo(
                id=template_data.get("id", ""),
                filename=template_data.get("filename", ""),
                display_name=template_data.get("display_name", {"ar": "", "en": ""}),
                description=template_data.get("description", {"ar": "", "en": ""}),
                placeholders=placeholders
            )
            
            # Check if the template file exists
            if os.path.exists(os.path.join(self.templates_dir, template_info.filename)):
                templates.append(template_info)
            else:
                print(f"Warning: Template file {template_info.filename} not found in {self.templates_dir}")
                
        return templates
    
    def get_template_by_id(self, template_id: str) -> Optional[TemplateInfo]:
        """Get a specific template by its ID"""
        templates = self.get_available_templates()
        for template in templates:
            if template.id == template_id:
                return template
        return None
    
    def get_template_by_filename(self, filename: str) -> Optional[TemplateInfo]:
        """Get a specific template by its filename"""
        templates = self.get_available_templates()
        for template in templates:
            if template.filename == filename:
                return template
        return None
    
    def _replace_text_in_paragraph(self, paragraph, placeholder_tag: str, value: str):
        """Replace placeholder in a paragraph while preserving formatting"""
        for run in paragraph.runs:
            if placeholder_tag in run.text:
                # Preserve the run's formatting while replacing text
                run.text = run.text.replace(placeholder_tag, str(value))
    
    def fill_template(self, template_id: str, values: Dict[str, Any]) -> Document:
        """Fill a template with provided values"""
        template = self.get_template_by_id(template_id)
        if not template:
            raise ValueError(f"Template with ID '{template_id}' not found")
        
        template_path = os.path.join(self.templates_dir, template.filename)
        doc = Document(template_path)
        
        # Replace placeholders in the document
        for placeholder in template.placeholders:
            if placeholder.id in values:
                value = values[placeholder.id]
                
                # Simply convert any value to string
                value_str = str(value)
                
                # Replace in paragraphs
                for paragraph in doc.paragraphs:
                    self._replace_text_in_paragraph(paragraph, placeholder.placeholder_tag, value_str)
                
                # Replace in tables
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                self._replace_text_in_paragraph(paragraph, placeholder.placeholder_tag, value_str)
                
                # Replace in headers and footers
                for section in doc.sections:
                    # Header
                    for paragraph in section.header.paragraphs:
                        self._replace_text_in_paragraph(paragraph, placeholder.placeholder_tag, value_str)
                    # Footer
                    for paragraph in section.footer.paragraphs:
                        self._replace_text_in_paragraph(paragraph, placeholder.placeholder_tag, value_str)
        
        return doc
    
    def save_document(self, doc: Document, output_path: str):
        """Save the document to the specified path"""
        # Create directory if it doesn't exist
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        doc.save(output_path)
        return output_path
    
    def convert_to_pdf(self, docx_path: str, pdf_path: str = None):
        """Convert a DOCX file to PDF"""
        if pdf_path is None:
            pdf_path = os.path.splitext(docx_path)[0] + ".pdf"
            
        try:
            if platform.system() == 'Windows' and self.pythoncom:
                # Initialize COM for the current thread
                self.pythoncom.CoInitialize()
            
            # Perform the conversion
            convert(docx_path, pdf_path)
            return pdf_path
            
        except Exception as e:
            print(f"Error converting to PDF: {str(e)}")
            raise
        
        finally:
            if platform.system() == 'Windows' and self.pythoncom:
                # Uninitialize COM
                self.pythoncom.CoUninitialize()
    
    def send_document_email(self, 
                            recipient_email: str,
                            subject: str,
                            body: str,
                            attachment_path: str,
                            smtp_server: str,
                            smtp_port: int,
                            smtp_username: str,
                            smtp_password: str):
        """Send an email with the document attached"""
        msg = MIMEMultipart()
        msg['From'] = smtp_username
        msg['To'] = recipient_email
        msg['Subject'] = subject
        
        msg.attach(MIMEText(body, 'plain'))
        
        with open(attachment_path, 'rb') as f:
            mime_type = 'pdf' if attachment_path.lower().endswith('.pdf') else 'docx'
            attachment = MIMEApplication(f.read(), _subtype=mime_type)
            attachment.add_header('Content-Disposition', 'attachment', filename=os.path.basename(attachment_path))
            msg.attach(attachment)
        
        with smtplib.SMTP(smtp_server, smtp_port) as server:
            server.starttls()
            server.login(smtp_username, smtp_password)
            server.send_message(msg)
    
    def generate_template_document(self, template_id: str, values: Dict[str, Any], output_format: str = "docx") -> str:
        """Generate a document from a template and return the path to the file"""
        # Generate temporary files
        temp_dir = tempfile.gettempdir()
        timestamp = str(int(os.path.getmtime(self.templates_config_path)) if os.path.exists(self.templates_config_path) else 0)
        temp_docx = os.path.join(temp_dir, f"temp_{template_id}_{timestamp}.docx")
        
        # Fill and save the document
        doc = self.fill_template(template_id, values)
        self.save_document(doc, temp_docx)
        
        # Convert to PDF if requested
        if output_format.lower() == "pdf":
            temp_pdf = os.path.splitext(temp_docx)[0] + ".pdf"
            self.convert_to_pdf(temp_docx, temp_pdf)
            return temp_pdf
        
        return temp_docx
    
    def cleanup_temp_file(self, file_path: str):
        """Clean up a temporary file"""
        try:
            if os.path.exists(file_path):
                os.remove(file_path)
                return True
        except Exception as e:
            print(f"Error cleaning up temporary file {file_path}: {str(e)}")
            return False